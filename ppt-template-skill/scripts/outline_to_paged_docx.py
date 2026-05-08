#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Convert a source outline into an editable paged Word planning document.

The Word document is intentionally plain and structured. Users can adjust
pagination, page titles, recommended layouts and page content, then convert it
back to content.json with paged_docx_to_content_json.py.
"""

from __future__ import annotations

import argparse
import re
import sys
import zipfile
from pathlib import Path
from typing import Dict, List
from xml.etree import ElementTree as ET

try:
    from docx import Document
    from docx.shared import Pt
except Exception as exc:  # pragma: no cover - user-facing dependency error
    raise SystemExit("Missing dependency: python-docx. Install with: pip install python-docx") from exc


NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def read_text(path: Path) -> str:
    if path.suffix.lower() == ".docx":
        with zipfile.ZipFile(path) as zf:
            xml = zf.read("word/document.xml")
        root = ET.fromstring(xml)
        lines: List[str] = []
        for para in root.findall(".//w:p", NS):
            parts = [node.text or "" for node in para.findall(".//w:t", NS)]
            text = "".join(parts).strip()
            if text:
                lines.append(text)
        return "\n".join(lines)
    for enc in ("utf-8", "utf-8-sig", "gb18030"):
        try:
            return path.read_text(encoding=enc)
        except UnicodeDecodeError:
            continue
    return path.read_text(errors="ignore")


def clean_lines(text: str) -> List[str]:
    return [re.sub(r"\s+", " ", line).strip() for line in text.splitlines() if line.strip()]


def is_heading(line: str) -> bool:
    return bool(
        re.match(r"^第?\s*\d+\s*[页章部分][：:\s]", line)
        or re.match(r"^#{1,4}\s+\S+", line)
        or re.match(r"^[一二三四五六七八九十]+[、.．]\s*\S+", line)
        or re.match(r"^\d+[、.．]\s+\S+", line)
    )


def strip_heading(line: str) -> str:
    line = re.sub(r"^#{1,4}\s+", "", line)
    line = re.sub(r"^第?\s*\d+\s*[页章部分][：:\s]*", "", line)
    line = re.sub(r"^[一二三四五六七八九十]+[、.．]\s*", "", line)
    line = re.sub(r"^\d+[、.．]\s+", "", line)
    return line.strip() or line


def split_into_pages(lines: List[str], max_body_lines: int = 7) -> List[Dict]:
    explicit: List[Dict] = []
    current = None
    for line in lines:
        m = re.match(r"^第\s*(\d+)\s*页[：:\s]*(.*)$", line)
        if m:
            if current:
                explicit.append(current)
            current = {"title": m.group(2).strip() or f"第{m.group(1)}页", "lines": []}
        elif current:
            current["lines"].append(line)
    if current:
        explicit.append(current)
    if explicit:
        return explicit

    pages: List[Dict] = []
    current = None
    for line in lines:
        if is_heading(line):
            if current:
                pages.append(current)
            current = {"title": strip_heading(line), "lines": []}
        elif current:
            current["lines"].append(line)
        else:
            current = {"title": line[:32], "lines": []}
    if current:
        pages.append(current)

    chunked: List[Dict] = []
    for page in pages:
        body = page["lines"] or [page["title"]]
        if len(body) <= max_body_lines:
            chunked.append(page)
            continue
        for idx in range(0, len(body), max_body_lines):
            suffix = "" if idx == 0 else f"（续{idx // max_body_lines + 1}）"
            chunked.append({"title": page["title"] + suffix, "lines": body[idx : idx + max_body_lines]})
    return chunked


def detect_variant(title: str, lines: List[str]) -> str:
    text = f"{title} " + " ".join(lines)
    digit_count = len(re.findall(r"\d", text))
    if any(k in text for k in ("核心观点", "金句", "结论", "判断", "范式", "必须")):
        return "T13_statement"
    if digit_count >= 8 or any(k in text for k in ("数据", "指标", "KPI", "NPS", "利润率", "增长", "%", "亿元")):
        return "T4d_metrics"
    if any(k in text for k in ("阶段", "路径", "演进", "时间线", "节奏", "计划", "排期")):
        return "T4e_timeline"
    if any(k in text for k in ("对比", "比较", "竞品", "风险", "应对", "矩阵", "过去", "现在")):
        return "T5_comparison"
    if any(k in text for k in ("表", "清单", "分工", "资源", "任务", "预算", "执行")):
        return "T7_table"
    if len(lines) >= 5:
        return "T4b_numbered_list"
    return "T4a_cards"


def field_lines_for_variant(variant: str, title: str, lines: List[str]) -> List[str]:
    if variant == "T13_statement":
        statement = next((x for x in lines if len(x) >= 12), title)
        support = [x for x in lines if x != statement][:4]
        out = [f"- statement：{statement}"]
        if support:
            out.append("- supporting_lines：")
            out.extend(f"  - {x}" for x in support)
        return out
    if variant == "T4d_metrics":
        return ["- metrics："] + [f"  - {line}" for line in lines[:6]]
    if variant == "T4e_timeline":
        return ["- stages："] + [f"  - {line}" for line in lines[:6]]
    if variant in ("T5_comparison", "T7_table"):
        return ["- rows："] + [f"  - {line}" for line in lines[:8]]
    return ["- items："] + [f"  - {line}" for line in lines[:8]]


def infer_sections(pages: List[Dict]) -> List[str]:
    titles = [p["title"] for p in pages if p["title"]]
    if not titles:
        return ["核心内容"]
    sections = []
    for title in titles:
        compact = re.sub(r"[：:｜|].*$", "", title).strip()
        if compact and compact not in sections:
            sections.append(compact)
        if len(sections) >= 4:
            break
    return sections or ["核心内容"]


def add_kv(doc: Document, key: str, value: str = "") -> None:
    doc.add_paragraph(f"{key}：{value}")


def write_docx(source: Path, pages: List[Dict], output: Path, title: str | None) -> None:
    deck_title = title or (pages[0]["title"] if pages else source.stem)
    sections = infer_sections(pages)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)

    doc.add_paragraph("CONFIRMATION_STATUS: pending")
    doc.add_paragraph("CONFIRMATION_REQUIRED: Review pagination, page budget, layout variants, and content rewrite boundaries before generating PPT.")

    doc.add_heading("PPT分页策划稿", level=0)
    add_kv(doc, "项目标题", deck_title)
    add_kv(doc, "来源文件", str(source))
    doc.add_paragraph("说明：可直接在本 Word 中调整页标题、页面类型、推荐排版和内容字段；确认后用 paged_docx_to_content_json.py 转为 content.json。")

    page_no = 1
    doc.add_heading(f"第 {page_no:02d} 页｜封面", level=1)
    add_kv(doc, "页面类型", "cover")
    add_kv(doc, "标题", deck_title)
    add_kv(doc, "副标题", "基于原始大纲自动分页生成")
    add_kv(doc, "分页理由", "保留模板封面页，仅替换标题相关文本。")
    add_kv(doc, "备注", "")

    page_no += 1
    doc.add_heading(f"第 {page_no:02d} 页｜目录", level=1)
    add_kv(doc, "页面类型", "toc")
    add_kv(doc, "标题", "目录")
    doc.add_paragraph("目录项：")
    for item in sections:
        doc.add_paragraph(f"- {item}")
    add_kv(doc, "分页理由", "用目录承接大纲结构，目录页仅填充文字，不改模板样式。")

    for page in pages:
        page_no += 1
        variant = detect_variant(page["title"], page["lines"])
        doc.add_heading(f"第 {page_no:02d} 页｜内容页", level=1)
        add_kv(doc, "页面类型", "content")
        add_kv(doc, "页面标题", page["title"])
        add_kv(doc, "推荐排版", variant)
        add_kv(doc, "排版理由", f"根据页面标题和正文特征匹配 {variant}，生成 PPT 时再渐进式加载对应排版脚本。")
        add_kv(doc, "分页理由", "同一主题内容聚合为一页；若字段过多，生成 PPT 前优先拆页，不删改大纲信息。")
        add_kv(doc, "容量提示", "请确认本页字段是否过多；如超过模板容量，可复制本页拆成续页。")
        doc.add_paragraph("内容字段：")
        for line in field_lines_for_variant(variant, page["title"], page["lines"]):
            doc.add_paragraph(line)
        doc.add_paragraph("来源：原始大纲")
        add_kv(doc, "备注", "")

    page_no += 1
    doc.add_heading(f"第 {page_no:02d} 页｜尾页", level=1)
    add_kv(doc, "页面类型", "end")
    add_kv(doc, "分页理由", "保留模板尾页原样，不替换任何文字。")

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> int:
    parser = argparse.ArgumentParser(description="Generate editable paged Word plan from an outline.")
    parser.add_argument("input", type=Path, help="Source outline: .docx, .txt or .md")
    parser.add_argument("output", type=Path, help="Output paged planning .docx")
    parser.add_argument("--title", help="Optional deck title")
    args = parser.parse_args()

    text = read_text(args.input)
    lines = clean_lines(text)
    if not lines:
        raise SystemExit(f"No usable text found in {args.input}")
    pages = split_into_pages(lines)
    write_docx(args.input, pages, args.output, args.title)
    print(f"Wrote {args.output} with {len(pages) + 3} planned slides.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
